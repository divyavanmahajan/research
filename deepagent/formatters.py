"""
Document formatters: converts markdown research output to MD, HTML, or DOCX.
"""

import re
from datetime import datetime
from pathlib import Path

CURRENT_DATE = datetime.now().strftime("%B %d, %Y")

# ──────────────────────────────────────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────────────────────────────────────

def save_document(content: str, fmt: str, path: str) -> None:
    """Dispatch to the correct formatter and write the output file."""
    fmt = fmt.lower().lstrip(".")
    Path(path).parent.mkdir(parents=True, exist_ok=True)

    title_match = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else "Enterprise Architecture Assessment"

    if fmt in ("md", "markdown"):
        save_as_markdown(content, path)
    elif fmt in ("html", "htm"):
        save_as_html(content, path, title)
    elif fmt in ("docx", "word"):
        save_as_docx(content, path, title)
    else:
        raise ValueError(f"Unsupported format '{fmt}'. Choose md, html, or docx.")


# ──────────────────────────────────────────────────────────────────────────────
# Markdown
# ──────────────────────────────────────────────────────────────────────────────

def save_as_markdown(content: str, path: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


# ──────────────────────────────────────────────────────────────────────────────
# HTML
# ──────────────────────────────────────────────────────────────────────────────

def save_as_html(content: str, path: str, title: str) -> None:
    import markdown2

    html_body = markdown2.markdown(
        content,
        extras=[
            "tables",
            "fenced-code-blocks",
            "header-ids",
            "strike",
            "footnotes",
        ],
    )

    toc_items = _build_html_toc(content)

    full_html = _HTML_TEMPLATE.format(
        title=title,
        date=CURRENT_DATE,
        toc_items=toc_items,
        content=html_body,
    )

    with open(path, "w", encoding="utf-8") as f:
        f.write(full_html)


def _build_html_toc(content: str) -> str:
    headings = re.findall(r"^(#{1,3})\s+(.+)$", content, re.MULTILINE)
    items = []
    for hashes, text in headings:
        level = len(hashes)
        clean = re.sub(r"[^\w\s-]", "", text.lower())
        anchor = re.sub(r"\s+", "-", clean).strip("-")
        indent = "  " * (level - 1)
        items.append(f"{indent}<li><a href='#{anchor}'>{text}</a></li>")
    return "\n".join(items)


_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Segoe UI',Calibri,'Helvetica Neue',Arial,sans-serif;font-size:14px;line-height:1.75;color:#1a2332;background:#eef0f4}}
.page-wrapper{{max-width:1060px;margin:0 auto;padding:24px}}

/* Cover */
.cover{{background:linear-gradient(160deg,#0a1628 0%,#0d2a4a 40%,#1a4a7a 100%);color:#fff;padding:56px 48px 48px;margin-bottom:6px;border-radius:4px 4px 0 0}}
.cover .badge{{font-size:10px;letter-spacing:2px;text-transform:uppercase;color:#7fb3d3;margin-bottom:16px}}
.cover h1{{font-size:2.2em;font-weight:700;line-height:1.25;color:#fff;margin-bottom:12px}}
.cover .subtitle{{font-size:1.05em;color:#a8c8e8;margin-bottom:32px}}
.cover .meta-row{{display:flex;gap:30px;font-size:12px;color:#7fb3d3;border-top:1px solid rgba(255,255,255,.15);padding-top:20px;flex-wrap:wrap}}
.cover .meta-row span strong{{color:#c8dff0}}

/* TOC */
.toc-section{{background:#fff;border-left:5px solid #1a4a7a;padding:24px 32px;margin-bottom:6px}}
.toc-section h2{{font-size:.8em;letter-spacing:1.5px;text-transform:uppercase;color:#1a4a7a;margin-bottom:14px}}
.toc-section ul{{padding-left:0;list-style:none;column-count:2;column-gap:30px}}
.toc-section li{{margin:5px 0}}
.toc-section a{{color:#1a4a7a;text-decoration:none;font-size:13px}}
.toc-section a:hover{{text-decoration:underline}}
.toc-section li li{{padding-left:15px;font-size:12px;color:#555}}

/* Content */
.content-body{{background:#fff;padding:36px 48px;margin-bottom:6px;border-radius:0 0 4px 4px}}
h1{{font-size:1.9em;color:#0a1628;border-bottom:3px solid #1a4a7a;padding-bottom:10px;margin:36px 0 18px}}
h1:first-child{{margin-top:0}}
h2{{font-size:1.35em;color:#0d2a4a;border-bottom:1px solid #c8dff0;padding-bottom:7px;margin:30px 0 14px}}
h3{{font-size:1.1em;color:#1a4a7a;margin:22px 0 10px}}
h4{{font-size:.95em;color:#2d6a9f;margin:16px 0 8px;font-weight:600}}
p{{margin:10px 0}}
strong{{color:#0d2a4a}}

/* Tables */
table{{width:100%;border-collapse:collapse;margin:18px 0;font-size:13px;box-shadow:0 1px 3px rgba(0,0,0,.08)}}
thead th{{background:#0d2a4a;color:#e8f4fd;padding:11px 14px;text-align:left;font-size:12px;letter-spacing:.5px;font-weight:600}}
tbody td{{padding:9px 14px;border-bottom:1px solid #e8edf3;vertical-align:top}}
tbody tr:nth-child(even){{background:#f5f8fb}}
tbody tr:hover{{background:#eaf1f8}}

/* Lists */
ul,ol{{padding-left:22px;margin:10px 0}}
li{{margin:5px 0}}

/* Blockquote */
blockquote{{border-left:4px solid #1a4a7a;background:#f0f5fa;padding:14px 20px;margin:16px 0;color:#2c3e50;font-style:italic}}

/* Code */
code{{background:#f0f2f5;border:1px solid #dde3ea;padding:2px 6px;border-radius:3px;font-family:'Consolas','Courier New',monospace;font-size:12px;color:#c0392b}}
pre{{background:#1e2a38;border-radius:4px;padding:18px;margin:16px 0;overflow-x:auto}}
pre code{{background:none;border:none;color:#e8f4fd;font-size:13px;padding:0}}

/* Risk severity colours */
td:has(strong){{}}
.sev-high{{color:#c0392b;font-weight:700}}
.sev-medium{{color:#d68910;font-weight:700}}
.sev-low{{color:#1e8449;font-weight:700}}

hr{{border:none;border-top:2px solid #e0e6ed;margin:30px 0}}

/* Footer */
.doc-footer{{text-align:center;color:#8899aa;font-size:11px;padding:20px;letter-spacing:.3px}}
</style>
</head>
<body>
<div class="page-wrapper">

  <div class="cover">
    <div class="badge">Enterprise Architecture Assessment</div>
    <h1>{title}</h1>
    <div class="subtitle">Strategic Technology Assessment &amp; Architecture Guidance</div>
    <div class="meta-row">
      <span><strong>Date:</strong> {date}</span>
      <span><strong>Version:</strong> 1.0</span>
      <span><strong>Status:</strong> Draft</span>
      <span><strong>Classification:</strong> Internal</span>
    </div>
  </div>

  <div class="toc-section">
    <h2>Table of Contents</h2>
    <ul>
{toc_items}
    </ul>
  </div>

  <div class="content-body">
    {content}
  </div>

  <div class="doc-footer">
    Enterprise Architecture Assessment &bull; {date} &bull; Version 1.0 &bull; Internal Use Only
  </div>

</div>
</body>
</html>"""


# ──────────────────────────────────────────────────────────────────────────────
# DOCX
# ──────────────────────────────────────────────────────────────────────────────

def save_as_docx(content: str, path: str, title: str) -> None:
    from docx import Document

    doc = Document()
    _setup_docx_styles(doc)
    _add_docx_cover(doc, title)
    _parse_markdown_to_docx(doc, content)
    _add_docx_footer(doc)
    doc.save(path)


def _setup_docx_styles(doc) -> None:
    from docx.shared import Pt, RGBColor

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for name, size, rgb in [
        ("Heading 1", 20, (0x0A, 0x16, 0x28)),
        ("Heading 2", 15, (0x0D, 0x2A, 0x4A)),
        ("Heading 3", 13, (0x1A, 0x4A, 0x7A)),
    ]:
        s = styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(*rgb)

    try:
        h4 = styles["Heading 4"]
        h4.font.name = "Calibri"
        h4.font.size = Pt(12)
        h4.font.bold = True
        h4.font.color.rgb = RGBColor(0x2D, 0x6A, 0x9F)
    except Exception:
        pass


def _add_docx_cover(doc, title: str) -> None:
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.2)
    section.right_margin = Cm(3.2)

    for _ in range(7):
        doc.add_paragraph()

    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = badge.add_run("ENTERPRISE ARCHITECTURE ASSESSMENT")
    br.font.size = Pt(9)
    br.font.color.rgb = RGBColor(0x2D, 0x6A, 0x9F)
    br.font.bold = True

    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_para.add_run(title)
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Strategic Technology Assessment & Architecture Guidance")
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor(0x1A, 0x4A, 0x7A)

    for _ in range(5):
        doc.add_paragraph()

    div = doc.add_paragraph("─" * 55)
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Date: {CURRENT_DATE}    |    Version: 1.0    |    Status: Draft    |    Classification: Internal"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

    doc.add_page_break()


def _add_docx_footer(doc) -> None:
    from docx.shared import Pt, RGBColor

    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.text = f"Enterprise Architecture Assessment  •  {CURRENT_DATE}  •  Internal Use Only"
    for run in para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)


def _parse_markdown_to_docx(doc, content: str) -> None:
    lines = content.split("\n")
    i = 0
    in_list = False

    while i < len(lines):
        line = lines[i]

        # Headings
        h_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if h_match:
            level = min(len(h_match.group(1)), 4)
            text = _strip_inline(h_match.group(2))
            doc.add_heading(text, level=level)
            in_list = False
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", line.strip()):
            doc.add_paragraph("─" * 55)
            i += 1
            continue

        # Table — collect consecutive pipe-rows
        if line.strip().startswith("|") and "|" in line:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not re.match(r"^\|[\s|:-]+\|$", row):  # skip separator
                    table_lines.append(row)
                i += 1
            if table_lines:
                _add_docx_table(doc, table_lines)
            in_list = False
            continue

        # Blockquote
        if line.startswith("> "):
            try:
                doc.add_paragraph(line[2:], style="Quote")
            except Exception:
                p = doc.add_paragraph()
                r = p.add_run(line[2:])
                r.italic = True
            i += 1
            continue

        # Unordered list
        ul = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if ul:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, ul.group(2))
            in_list = True
            i += 1
            continue

        # Ordered list
        ol = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if ol:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, ol.group(2))
            in_list = True
            i += 1
            continue

        # Empty line
        if not line.strip():
            in_list = False
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, line)
        in_list = False
        i += 1


def _add_docx_table(doc, rows_text: list) -> None:
    from docx.shared import RGBColor

    rows = []
    for line in rows_text:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row_cells in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx < max_cols:
                cell = row.cells[c_idx]
                clean = _strip_inline(cell_text)
                cell.text = clean
                if r_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0x0D, 0x2A, 0x4A)


def _add_inline_runs(paragraph, text: str) -> None:
    from docx.shared import Pt

    # Split on bold, italic, code, and links
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))", text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif re.match(r"\[[^\]]+\]\([^)]+\)", part):
            link_text = re.match(r"\[([^\]]+)\]", part).group(1)
            run = paragraph.add_run(link_text)
            run.underline = True
        else:
            paragraph.add_run(part)


def _strip_inline(text: str) -> str:
    """Remove markdown inline markers, returning plain text."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()
